## firma/cerrar_firma.py
from redmine.redmine_client import actualizar_estado_issue
from firma.firma_mailer import enviar_docx_final
from app import db
import logging
import os
from docx import Document
from dotenv import load_dotenv

# Cargar configuración
load_dotenv()
ENVIAR_FIRMADO = os.getenv("ENVIAR_DOC_FIRMADO", "false").lower() == "true"

def cerrar_rechazado(documento):
    try:
        logging.info(f"cerrar_firma - Iniciando cierre por rechazo para documento ID {documento.id}")
        actualizar_estado_issue(documento.issue_id, nuevo_estado_id=15)  # Estado 'Firma Rechazada'
        documento.procesado = True
        db.session.commit()
        logging.info(f"cerrar_firma - Documento {documento.id} marcado como rechazado y flujo cerrado.")

    except Exception as e:
        logging.error(f"cerrar_firma - ❌ Error al cerrar documento rechazado {documento.id}: {e}")

def cerrar_aprobado(documento):
    try:
        ruta_docx = documento.path
        logging.info(f"cerrar_firma - Iniciando cierre por aprobación para documento ID {documento.id}")

        # 1. Estampar el documento
        logging.info(f"cerrar_firma - Estampando documento en {ruta_docx}")
        doc = Document(ruta_docx)
        doc.add_paragraph("\n---\nDocumento aprobado electrónicamente.")
        for firma in documento.firmas:
            if firma.estado == "aceptado":
                texto_firma = (
                    f"Firmado por: {firma.nombre}\n"
                    f"RUT: {firma.rut}\n"
                    f"Fecha: {firma.fecha_firma.strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"Estado: ACEPTADO"
                )
                doc.add_paragraph(texto_firma)
        doc.save(ruta_docx)
        logging.info(f"cerrar_firma - Documento estampado correctamente.")

        # 2. Cambiar estado en Redmine
        actualizar_estado_issue(documento.issue_id, nuevo_estado_id=14)
        logging.info(f"cerrar_firma - Estado del issue {documento.issue_id} actualizado a 14 (Firmado)")

        # 3. Enviar por correo si está habilitado
        if ENVIAR_FIRMADO:
            logging.info(f"cerrar_firma - Enviando documento firmado a {documento.responsable_email}")
            enviar_docx_final(documento, ruta_docx)

        # 4. Marcar como procesado
        documento.procesado = True
        db.session.commit()
        logging.info(f"cerrar_firma - Documento {documento.id} firmado correctamente y flujo completado.")

    except Exception as e:
        logging.error(f"cerrar_firma - ❌ Error al cerrar documento aprobado {documento.id}: {e}")
