## estampar_firmas.py
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utils.logging_config import configurar_logging
import logging
configurar_logging()
logger = logging.getLogger(__name__)

def estampar_firmas(issue_id, nombre_documento, path_documento, firmas_requeridas, token_documento):
    try:
        logger.info(f"issue_id: {issue_id}")
        logger.info(f"nombre_documento: {nombre_documento}")
        logger.info(f"path_documento: {path_documento}")
        logger.info(f"token_documento: {token_documento}")

        ruta_completa = os.path.join(path_documento, nombre_documento)
        logger.info(f"Abriendo documento para estampado: {ruta_completa}")

        if not os.path.isfile(ruta_completa):
            logger.error(f"Documento no encontrado: {ruta_completa}")
            return False

        doc = Document(ruta_completa)

        # Organizar firmas por tipo
        firmante = next((f for f in firmas_requeridas if f.tipo == 'firmante'), None)
        responsable = next((f for f in firmas_requeridas if f.tipo == 'responsable'), None)

        for section in doc.sections:
            footer = section.footer

            # Calcular el ancho útil entre márgenes
            ancho_total = section.page_width - section.left_margin - section.right_margin

            # Crear tabla con una fila y dos columnas
            table = footer.add_table(rows=1, cols=2, width=ancho_total)
            table.allow_autofit = False
            table.columns[0].width = ancho_total // 2
            table.columns[1].width = ancho_total // 2

            # Eliminar bordes de la tabla
            tbl = table._tbl
            tblPr = tbl.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
            borders = OxmlElement('w:tblBorders')
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                borders.append(border)
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            tblPr.append(borders)
            if tbl.find(qn('w:tblPr')) is None:
                tbl.insert(0, tblPr)

            row_cells = table.rows[0].cells

            # Firmante a la izquierda
            if firmante:
                p = row_cells[0].paragraphs[0]
                p.alignment = 0  # LEFT
                run = p.add_run(f"Trabajador: {firmante.nombre} {firmante.rut} {firmante.fecha_firma}")
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Responsable a la derecha
            if responsable:
                p = row_cells[1].paragraphs[0]
                p.alignment = 2  # RIGHT
                run = p.add_run(f"Empresa: {responsable.nombre} {responsable.rut} {responsable.fecha_firma}")
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Insertar sección final con código de validación
        doc.add_paragraph()  # Separación

        p = doc.add_paragraph()
        run = p.add_run("Este documento fue firmado electrónicamente.")
        run.font.size = Pt(10)
        run.bold = True

        p = doc.add_paragraph()
        run = p.add_run(f"Código de validación: {token_documento}")
        run.font.size = Pt(10)
        run.italic = True

        p = doc.add_paragraph()
        run = p.add_run("Puede verificar la autenticidad en:")
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run("https://condominium.eproc-chile.cl/validar")
        run.font.size = Pt(10)
        run.underline = True
        run.font.color.rgb = RGBColor(0, 0, 255)

        # Guardar documento modificado
        doc.save(ruta_completa)
        logger.info(f"Estampado completado exitosamente para issue_id: {issue_id}")
        return True

    except Exception as e:
        logger.error(f"Error al estampar firmas: {e}", exc_info=True)
        return False
